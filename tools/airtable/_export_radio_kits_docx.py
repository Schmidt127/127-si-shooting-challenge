#!/usr/bin/env python3
"""Export radio media kit markdown files to styled Word documents."""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from media_paths import radio_root

MD_NAME = "RADIO-MEDIA-KIT.md"
DOCX_NAME = "RADIO-MEDIA-KIT.docx"

BRAND = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x74, 0xB5)
MUTED = RGBColor(0x59, 0x59, 0x59)


def set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_formatted_runs(paragraph, text: str, *, base_size: int = 11, base_bold: bool = False):
    """Parse **bold** segments into Word runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, size=base_size)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, bold=base_bold, size=base_size)


def add_horizontal_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, space_before, space_after, color in (
        ("Heading 1", 16, 18, 8, BRAND),
        ("Heading 2", 13, 14, 6, BRAND),
        ("Heading 3", 12, 10, 4, ACCENT),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.replace("# ", "").strip())
    set_run_font(run, bold=True, size=22, color=BRAND)
    p.paragraph_format.space_after = Pt(4)

    rule = doc.add_paragraph()
    add_horizontal_rule(rule)
    rule.paragraph_format.space_after = Pt(10)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip("* "))
    set_run_font(run, italic=True, size=10, color=MUTED)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    clean = text.lstrip("#").strip()
    doc.add_heading(clean, level=min(level, 3))


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_formatted_runs(p, text)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    add_formatted_runs(p, text.lstrip("-* ").strip())
    p.paragraph_format.space_after = Pt(2)


def add_blockquote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    add_formatted_runs(p, text.lstrip("> ").strip(), base_size=11)
    for run in p.runs:
        run.italic = True
        run.font.color.rgb = MUTED


def add_metadata_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_formatted_runs(p, text, base_size=10)


def add_footer_note(doc: Document) -> None:
    doc.add_paragraph()
    rule = doc.add_paragraph()
    add_horizontal_rule(rule)
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("127 Sports Intensity – Shooting Challenge (2026)")
    set_run_font(run, bold=True, size=10, color=BRAND)

    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link = url.add_run("www.fairfieldbasketballclub.com/leaderboard")
    set_run_font(link, size=10, color=ACCENT)


def markdown_to_docx(md_text: str, doc: Document) -> None:
    configure_document(doc)
    lines = md_text.splitlines()
    i = 0
    title_done = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# ") and not title_done:
            add_title_block(doc, line)
            title_done = True
            i += 1
            if i < len(lines) and lines[i].startswith("*Generated"):
                add_subtitle(doc, lines[i])
                i += 1
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:], 3)
            i += 1
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:], 2)
            i += 1
            continue

        if line.strip() == "---":
            rule = doc.add_paragraph()
            add_horizontal_rule(rule)
            rule.paragraph_format.space_before = Pt(6)
            rule.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        if line.startswith("> "):
            add_blockquote(doc, line)
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            add_bullet(doc, line)
            i += 1
            continue

        if line.startswith("**Market ID:**") or line.startswith("**Communities:**") or line.startswith(
            "**Athletes in kit:**"
        ):
            add_metadata_line(doc, line)
            i += 1
            continue

        add_body_paragraph(doc, line)
        i += 1

    add_footer_note(doc)


def export_kit(md_path: Path) -> Path:
    out_path = md_path.with_name(DOCX_NAME)
    doc = Document()
    md_text = md_path.read_text(encoding="utf-8")
    markdown_to_docx(md_text, doc)
    doc.save(out_path)
    return out_path


def main() -> None:
    root = radio_root()
    kits = sorted(root.glob(f"*/{MD_NAME}"))
    if not kits:
        raise SystemExit(f"No {MD_NAME} files found under {root}")

    exported: list[tuple[str, Path]] = []
    for md_path in kits:
        out = export_kit(md_path)
        exported.append((md_path.parent.name, out))
        print(f"  {out.relative_to(root)}")

    summary = root / "DOCX-EXPORT-SUMMARY.md"
    stamp = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M %Z")
    lines = [
        "# Radio Media Kit — Word Export Summary",
        "",
        f"Generated: {stamp}",
        "",
        f"Exported **{len(exported)}** kits to `{DOCX_NAME}` alongside each markdown source.",
        "",
        "| Market folder | Word file |",
        "|---------------|-----------|",
    ]
    for folder, path in exported:
        lines.append(f"| `{folder}` | `{path.name}` |")
    lines.extend(
        [
            "",
            "Open in Microsoft Word for final review before attaching to station emails.",
            "",
        ]
    )
    summary.write_text("\n".join(lines), encoding="utf-8", newline="\n")
    print(f"\nExported {len(exported)} kits -> {root}")
    print(f"Summary -> {summary}")


if __name__ == "__main__":
    main()
